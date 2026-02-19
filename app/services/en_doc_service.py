"""
영문 재무제표 파일(Word/PDF) → EnDocument 변환 서비스.
핵심 목표: 각 Note 섹션의 원문 텍스트를 최대한 온전히 보존.
숫자 추출/구조화 금지 — 그것은 LLM의 역할.
"""
import logging
import re
from pathlib import Path

from app.models.en_doc_model import DocFormat, EnDocument, EnNote

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────
# Note 섹션 분리 공통 패턴
# PDF: "1. General Information", "15. Tax Expense and Deferred Tax"
# (하위항목 "1.1", "2.2.1" 은 소수점 포함이므로 자동 제외)
# ─────────────────────────────────────────────────────
_NOTE_RE = re.compile(r"^\s*(\d+)\.\s+([A-Z][^\n]{0,120})$")

# 페이지 헤더/푸터 노이즈 패턴 (PDF 특화)
_PDF_HEADER_RE = re.compile(
    r"(GOWOONSESANG.{0,80}Notes to.{0,80}December 31)",
    re.IGNORECASE,
)
_PAGE_NUM_RE = re.compile(r"^\d{1,3}$")

# Word 파서에서 Note 제목으로 인식할 스타일 이름들
_WORD_TITLE_STYLES = {"ABCTitle", "Heading 1", "heading 1", "Title"}


# ─────────────────────────────────────────────────────
# 공개 API
# ─────────────────────────────────────────────────────

async def parse_en_file(file_path: Path) -> EnDocument:
    """
    업로드된 영문 파일을 파싱하여 EnDocument 반환.
    포맷은 확장자 우선으로 자동 감지.
    """
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return await _parse_word(file_path)
    elif suffix == ".pdf":
        return await _parse_pdf(file_path)
    else:
        # magic bytes fallback
        with open(file_path, "rb") as f:
            header = f.read(8)
        if header[:4] == b"PK\x03\x04":  # ZIP → docx
            return await _parse_word(file_path)
        elif header[:4] == b"%PDF":
            return await _parse_pdf(file_path)
        raise ValueError(
            f"지원하지 않는 파일 형식입니다: {file_path.name}\n"
            ".docx 또는 .pdf 파일만 허용됩니다."
        )


# ─────────────────────────────────────────────────────
# Word 파서
# ─────────────────────────────────────────────────────

async def _parse_word(file_path: Path) -> EnDocument:
    """
    python-docx 기반 Word(.docx) 파서.
    ABCTitle 또는 Heading 1 스타일 단락을 Note 제목으로 인식.
    단락 + 테이블을 DOM 순서로 순회하여 raw_text 구성.
    """
    import docx
    from docx.oxml.ns import qn

    doc = docx.Document(str(file_path))

    # DOM 순서대로 단락과 테이블을 함께 순회
    # (doc.paragraphs는 테이블 내 단락 포함 불가 → body 직접 순회)
    body_elements = list(doc.element.body)
    para_elements = {id(p._element): p for p in doc.paragraphs}
    table_elements = {id(t._element): t for t in doc.tables}

    # Note 섹션 분리
    sections: list[tuple[str, str, list[str]]] = []  # (note_num, title, lines)
    current_num: str | None = None
    current_title: str = ""
    current_lines: list[str] = []
    note_counter = 0

    def flush():
        nonlocal current_num, current_title, current_lines
        if current_num is not None:
            sections.append((current_num, current_title, list(current_lines)))
        current_num = None
        current_title = ""
        current_lines = []

    for elem in body_elements:
        elem_id = id(elem)

        # 단락 처리
        if elem_id in para_elements:
            para = para_elements[elem_id]
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            # 페이지 번호 단락 제외
            if _PAGE_NUM_RE.match(text):
                continue

            # Note 제목 감지 (ABCTitle 스타일)
            if style_name in _WORD_TITLE_STYLES:
                flush()
                note_counter += 1
                current_num = str(note_counter)
                current_title = text
                current_lines = [text]
            else:
                if current_num is not None:
                    current_lines.append(text)

        # 테이블 처리
        elif elem_id in table_elements:
            if current_num is None:
                continue
            table = table_elements[elem_id]
            table_text = _table_to_text(table)
            if table_text:
                current_lines.append(table_text)

    flush()

    full_text = "\n\n".join(
        "\n".join(lines) for _, _, lines in sections
    )
    notes = _sections_to_notes(sections, DocFormat.WORD)

    logger.info("Word 파싱 완료: Note %d개 (파일: %s)", len(notes), file_path.name)

    if len(notes) < 3:
        logger.warning("Note 3개 미만 — fallback 모드로 전환")
        return EnDocument(
            filename=file_path.name,
            format=DocFormat.WORD,
            notes=[],
            full_raw_text=_word_full_text(doc),
        )

    return EnDocument(
        filename=file_path.name,
        format=DocFormat.WORD,
        notes=notes,
        full_raw_text=full_text,
    )


def _word_full_text(doc) -> str:
    """Word 문서 전체 텍스트 추출 (fallback용)."""
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not _PAGE_NUM_RE.match(text):
            parts.append(text)
    return "\n".join(parts)


def _table_to_text(table) -> str:
    """
    python-docx Table 객체를 탭/줄바꿈 구분 텍스트로 변환.
    병합 셀(빈 문자열)도 탭으로 위치 보존.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append("\t".join(cells))
    return "\n".join(rows)


# ─────────────────────────────────────────────────────
# PDF 파서
# ─────────────────────────────────────────────────────

async def _parse_pdf(file_path: Path) -> EnDocument:
    """
    pdfplumber 기반 PDF 파서.
    페이지별 텍스트 + 테이블 추출 후 Note 섹션 분리.
    """
    import pdfplumber

    page_texts: list[tuple[int, str]] = []  # (page_num, text)

    with pdfplumber.open(str(file_path)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # 1) 일반 텍스트 추출
            raw = page.extract_text(x_tolerance=3, y_tolerance=3) or ""

            # 2) 테이블이 있는 페이지는 별도 추출 후 병합
            tables = page.extract_tables()
            table_texts = [_pdf_table_to_text(t) for t in (tables or []) if t]

            # 텍스트 정리 (페이지 헤더 노이즈 제거)
            cleaned = _clean_pdf_page(raw)

            # 테이블 텍스트를 본문 뒤에 추가
            if table_texts:
                cleaned = cleaned + "\n" + "\n".join(table_texts)

            if cleaned.strip():
                page_texts.append((page_num, cleaned))

    full_text = "\n\n".join(text for _, text in page_texts)

    # Note 섹션 분리
    sections = _split_pdf_into_sections(page_texts)
    notes = _sections_to_notes(sections, DocFormat.PDF)

    logger.info("PDF 파싱 완료: Note %d개 (파일: %s)", len(notes), file_path.name)

    if len(notes) < 3:
        logger.warning("Note 3개 미만 — fallback 모드로 전환")
        return EnDocument(
            filename=file_path.name,
            format=DocFormat.PDF,
            notes=[],
            full_raw_text=full_text,
        )

    return EnDocument(
        filename=file_path.name,
        format=DocFormat.PDF,
        notes=notes,
        full_raw_text=full_text,
    )


def _clean_pdf_page(text: str) -> str:
    """
    PDF 페이지 텍스트에서 헤더/푸터 노이즈 제거.
    """
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # 페이지 번호 단독 행 제거
        if _PAGE_NUM_RE.match(stripped):
            continue
        # 반복 헤더 제거 (회사명 + "Notes to ..." + 날짜)
        if _PDF_HEADER_RE.search(stripped):
            continue
        if re.match(r"December 31,\s*\d{4}", stripped, re.IGNORECASE):
            continue
        # 하이픈으로 끊긴 단어 복원
        if cleaned and cleaned[-1].endswith("-"):
            cleaned[-1] = cleaned[-1][:-1] + stripped
        else:
            cleaned.append(stripped)
    return "\n".join(cleaned)


def _pdf_table_to_text(table: list[list]) -> str:
    """
    pdfplumber extract_tables() 결과를 탭/줄바꿈 텍스트로 변환.
    None 셀은 빈 문자열로 처리.
    """
    rows = []
    for row in table:
        cells = [str(cell).replace("\n", " ").strip() if cell is not None else "" for cell in row]
        rows.append("\t".join(cells))
    return "\n".join(rows)


def _split_pdf_into_sections(
    page_texts: list[tuple[int, str]],
) -> list[tuple[str, str, list[str]]]:
    """
    PDF 전체 텍스트에서 Note 섹션 분리.
    반환: [(note_num, title, lines), ...]
    """
    sections: list[tuple[str, str, list[str]]] = []
    current_num: str | None = None
    current_title: str = ""
    current_lines: list[str] = []
    current_start_page: int = 1

    def flush():
        nonlocal current_num
        if current_num is not None:
            sections.append((current_num, current_title, list(current_lines)))
        current_num = None

    for page_num, text in page_texts:
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            # Note 제목 패턴 감지 (소수점 포함 번호 제외: "1.1", "2.2.1" 등)
            m = _NOTE_RE.match(stripped)
            if m and not re.match(r"^\d+\.\d+", stripped):
                flush()
                current_num = m.group(1)
                current_title = m.group(2).strip()
                current_lines = [stripped]
                current_start_page = page_num
            elif current_num is not None:
                current_lines.append(stripped)

    flush()
    return sections


# ─────────────────────────────────────────────────────
# 공통 헬퍼
# ─────────────────────────────────────────────────────

def _sections_to_notes(
    sections: list[tuple[str, str, list[str]]],
    fmt: DocFormat,
) -> list[EnNote]:
    """
    sections 리스트 → EnNote 리스트 변환.
    """
    notes = []
    for note_num, title, lines in sections:
        raw_text = "\n".join(lines)
        notes.append(
            EnNote(
                note_number=note_num,
                note_title=title,
                raw_text=raw_text,
                source_format=fmt,
            )
        )
    return notes
